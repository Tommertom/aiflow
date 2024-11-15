# aiflow.py
import json

from openai import OpenAI, OpenAIError
from typing import Optional, Callable, List, Dict, Union
from enum import Enum
from IPython.display import HTML
from docx import Document
import urllib.request
from pathlib import Path
import os
from IPython.display import Markdown, display
import markdown
from pydantic import BaseModel


# chroma helper that converts a query result to a string, so we can use it in the class
def chroma_query_result_to_text(obj):
    documents = obj.get("documents")
    if documents:
        concatenated_string = "".join(["\n".join(doc) for doc in documents])
        return concatenated_string
    else:
        return ""


# chroma helper that converts the query to a list
def chroma_query_to_list(result):
    ids = result["ids"][0]
    metadatas = result["metadatas"][0]
    distances = result["distances"][0]

    # combine each n-th item from ids, metadatas and distances and put it in an object. Append to a result list and return
    if len(ids) != len(metadatas) or len(ids) != len(distances):
        raise ValueError("Lengths of ids, metadatas, and distances must be the same")

    result_list = []
    for i in range(len(ids)):
        # Create an object with id, metadata, and distance
        item = {"id": ids[i], "metadata": metadatas[i], "distance": distances[i]}
        # Append the object to the result list
        result_list.append(item)

    return result_list


class Model(Enum):
    GPT_4 = "gpt-4"
    GPT_4O = "gpt-4o"
    GPT_4O_MINI = "gpt-4o-mini"


class AIFlow:

    def __init__(
        self,
        api_key,
        model=Model.GPT_4O_MINI,
        temperature=0,
        max_tokens=150,
    ):
        self.client = OpenAI(api_key=api_key)
        self.model = model.value
        self.max_tokens = max_tokens
        self.temperature = temperature
        self.json_mode = False
        self.completion_tokens = 0
        self.prompt_tokens = 0
        self.total_tokens = 0

        self.chat_messages = []
        self.context_map = {}
        self.images_map = {}
        self.audio_map = {}

        self.default_folder_for_output = ""
        self.verbose = True
        self.latest_state_filename = ""
        self.save_state_per_step = False

    # model configs
    def set_temperature(self, temperature: int = 0) -> "AIFlow":
        """
        Set the temperature for the model.

        :param temperature: Temperature value
        :return: self
        """
        self.temperature = temperature
        return self
        self.temperature = temperature
        return self

    def set_model(self, model: Union[Model, str] = Model.GPT_4) -> "AIFlow":
        """
        Set the model to be used.

        :param model: Model name
        :return: self
        """
        self.model = model.value
        return self

    def set_max_tokens(self, max_tokens: int = 150) -> "AIFlow":
        """
        Set the maximum number of tokens.

        :param max_tokens: Maximum tokens value
        :return: self
        """
        self.max_tokens = max_tokens
        return self

    def set_json_output(self, json_mode: bool = False) -> "AIFlow":
        """
        Set the output format to JSON.

        :param json_mode: Boolean flag for JSON mode
        :return: self
        """
        self.json_mode = json_mode
        return self

    def display_model_config(self) -> "AIFlow":
        """
        Display the current model configuration.

        :return: self
        """
        print(f"Model: {self.model}")
        print(f"Max Tokens: {self.max_tokens}")
        print(f"Temperature: {self.temperature}")
        return self

    def get_token_usage(self) -> Dict[str, int]:
        """
        Get the token usage statistics.

        :return: Dictionary with token usage statistics
        """
        return {
            "completion_tokens": self.completion_tokens,
            "prompt_tokens": self.prompt_tokens,
            "total_tokens": self.total_tokens,
        }

    # other config
    def set_output_folder(self, folder: str = "") -> "AIFlow":
        """
        Set the default folder for output.

        :param folder: Folder path
        :return: self
        """
        self.default_folder_for_output = folder
        if folder != "":
            os.makedirs(self.default_folder_for_output, exist_ok=True)
        return self

    def set_verbose(self, level: bool = True) -> "AIFlow":
        """
        Set the verbosity level.

        :param level: Boolean flag for verbosity
        :return: self
        """
        self.verbose = level
        return self

    def set_step_save(self, step: bool = False) -> "AIFlow":
        """
        Enable or disable saving state per step.

        :param step: Boolean flag for step save
        :return: self
        """
        self.save_state_per_step = step
        return self

    #
    # Saving state
    #

    def save_internal_state(self, filename: str = "") -> "AIFlow":
        """
        Save the internal state to a file.

        :param filename: Name of the file to save the state
        :return: self
        """
        if filename == "" and self.latest_state_filename == "":
            print("Error - no state filename provided")
            return self

        if filename == "":
            filename = self.latest_state_filename

        self.latest_state_filename = filename

        # Create a copy of the current state
        state_to_save = self.__dict__.copy()

        # Remove objects that cannot be serialized
        state_to_save.pop("client", None)

        # Save the state to a JSON file
        with open(filename, "w") as f:
            json.dump(state_to_save, f, indent=4)

        return self

    def load_internal_state(self, filename: str = "state.json") -> "AIFlow":
        """
        Load the internal state from a file.

        :param filename: Name of the file to load the state from
        :return: self
        """
        self.latest_state_filename = filename
        try:
            with open(filename, "r") as f:
                state = json.load(f)

            self.__dict__.update(state)
        except FileNotFoundError:
            print(f"File '{filename}' not found.")
        return self

    #
    # Some debugging tools
    #
    def display_internal_data(self) -> "AIFlow":
        """
        Display internal data for debugging.

        :return: self
        """
        print("Chat Messages:")
        print(json.dumps(self.chat_messages, indent=4))
        print("\nContext Map:")
        print(json.dumps(self.context_map, indent=4))
        print("\nImages Map:")
        print(json.dumps(self.images_map, indent=4))
        print("\nAudio Map:")
        print(json.dumps(self.audio_map, indent=4))
        return self

    def clear_internal_data(self) -> "AIFlow":
        """
        Clear internal data.

        :return: self
        """
        self.chat_messages = []
        self.context_map = {}
        self.images_map = {}
        self.audio_map = {}
        return self

    # function to run another function that may return something or nothing - this to support running code in the chain
    def execute_function(
        self, func: Callable[[], str] = lambda: "", label: str = ""
    ) -> "AIFlow":
        """
        Run a function that may return something or nothing.

        :param func: Function to run
        :param label: Label for the context
        :return: self
        """
        return self

    #
    # Chat methods
    #
    def pretty_print_messages(self) -> "AIFlow":
        """
        Pretty print chat messages.

        :return: self
        """
        for message in self.chat_messages:
            role = message["role"]
            content = message["content"]
            print(f"{role}:")
            print(content)
            print()
        return self

    def pretty_print_messages_to_file(
        self, file_name: str = "output.txt", html: bool = True
    ) -> "AIFlow":
        """
        Pretty print chat messages to a file.

        :param file_name: Name of the file to save the messages
        :param html: Whether to return HTML for downloading the file
        :return: self
        """
        with open(file_name, "w") as file:
            for message in self.chat_messages:
                role = message["role"]
                content = message["content"]
                file.write(f"{role}:\n")
                file.write(content + "\n\n")
        if html:
            return HTML(
                f'<a href="{file_name}" download>Click here to download the pretty-printed messages</a>'
            )
        return self

    def set_system_prompt(self, prompt: str = "") -> "AIFlow":
        """
        Set the system prompt.

        :param prompt: System prompt
        :return: self
        """
        # Remove existing "system" role message if it exists
        self.chat_messages = [
            msg for msg in self.chat_messages if msg.get("role") != "system"
        ]

        if self.verbose:
            print(prompt)

        if self.save_state_per_step:
            self.save_internal_state()

        prompt = self.replace_tags_with_content(prompt)

        # Insert new system message at the beginning of the list
        self.chat_messages.insert(0, {"role": "system", "content": prompt})
        return self

    def add_user_chat(self, prompt: str, label: str = "latest") -> "AIFlow":
        """
        Add a user chat message and get a response.

        :param prompt: User chat message
        :param label: Label for the context
        :return: self
        """
        if self.verbose:
            print(prompt)

        prompt = self.replace_tags_with_content(prompt)

        self.context_map["latest_prompt"] = prompt
        self.chat_messages.append({"role": "user", "content": prompt})

        response = self.call_openai_chat_api(messages=self.chat_messages)

        self.chat_messages.append({"role": "assistant", "content": response})
        self.context_map["latest"] = response
        self.context_map[label] = response

        if self.verbose:
            print(response)

        if self.save_state_per_step:
            self.save_internal_state()

        return self

    def filter_messages(
        self, func: Optional[Callable[[List[Dict[str, str]]], List[Dict[str, str]]]]
    ) -> "AIFlow":
        """
        Filter chat messages using a function.

        :param func: Function to filter messages
        :return: self
        """
        if func is not None:
            self.chat_messages = func(self.chat_messages)
        return self

    def reduce_messages_to_text(
        self, func: Optional[Callable[[List[Dict[str, str]]], str]]
    ) -> "AIFlow":
        """
        Reduce chat messages to text using a function.

        :param func: Function to reduce messages
        :return: self
        """
        if func is not None:
            self.context_map["latest"] = func(self.chat_messages)
            print(self.context_map["latest"])
        return self

    #
    # Simple completion
    #
    def generate_completion(self, prompt: str, label: str = "latest") -> "AIFlow":
        if self.verbose:
            print(prompt)
            print()

        prompt = self.replace_tags_with_content(prompt)
        self.context_map["latest_prompt"] = prompt
        messages = [{"role": "user", "content": prompt}]
        response = self.call_openai_chat_api(messages=messages)
        self._update_context(response, label)

        if self.verbose:
            print(response)
            print()

        if self.save_state_per_step:
            self.save_internal_state()

        return self

    #
    # Simple completion
    #
    def generate_json_completion(
        self, prompt: str, label: str = "latest", schema=BaseModel
    ) -> "AIFlow":

        if self.verbose:
            print(prompt)
            print()

        prompt = self.replace_tags_with_content(prompt)
        self.context_map["latest_prompt"] = prompt
        messages = [{"role": "user", "content": prompt}]
        response = self.call_openai_parse_api(messages=messages, schema=schema)
        self._update_context(response.json(), label)

        if self.verbose:
            print(response)
            print()

        if self.save_state_per_step:
            self.save_internal_state()

        return self

    def _update_context(self, response: str, label: str) -> None:
        self.context_map["latest"] = response
        self.context_map[label] = response

    #
    # OpenAI caller - completions
    #
    def call_openai_chat_api(self, messages: List[Dict[str, str]] = []) -> str:
        """
        Call the OpenAI chat API with the given messages.

        :param messages: List of messages
        :return: Response from the API
        """
        params = {
            "model": self.model,
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": messages,
        }

        # Conditionally add response_format
        if self.json_mode:
            params["response_format"] = {"type": "json_object"}

        try:
            completion = self.client.chat.completions.create(**params)
            self.update_token_usage(completion.usage)
            return completion.choices[0].message.content
        except OpenAIError as e:
            print(f"OpenAI API error: {e}")
            return "An error occurred with the OpenAI API."
        except Exception as e:
            print(f"Unexpected error: {e}")
            return "An unexpected error occurred."

    #
    # OpenAI caller - json_schema
    #
    def call_openai_parse_api(
        self, messages: List[Dict[str, str]] = [], schema=BaseModel
    ) -> str:
        """
        Call the OpenAI chat API with the given messages and the schema to generate JONS

        :param messages: List of messages
        :param schema: Schema based on Pydantic BaseModel
        :return: Response from the API
        """
        params = {
            "model": self.model,
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": messages,
            "response_format": schema,
        }

        try:
            completion = self.client.beta.chat.completions.parse(**params)
            self.update_token_usage(completion.usage)
            return completion.choices[0].message.parsed
        except OpenAIError as e:
            print(f"OpenAI API error: {e}")
            return "An error occurred with the OpenAI API."
        except Exception as e:
            print(f"Unexpected error: {e}")
            return "An unexpected error occurred."

    #
    # Completing context in prompts
    #
    def replace_tags_with_content(self, input_string: str = "") -> str:
        """
        Replace tags in the input string with context content.

        :param input_string: Input string with tags
        :return: String with tags replaced by context content
        """
        while True:
            previous_string = input_string
            for key, value in self.context_map.items():
                input_string = input_string.replace(f"[{key}]", str(value))

            # If no replacements were made in this iteration, break the loop
            if input_string == previous_string:
                break

        return input_string

    #
    # Context functions
    #
    def copy_latest_to(self, label: str = "latest") -> "AIFlow":
        """
        Copy the latest context to a specified label.

        :param label: Label for the context
        :return: self
        """
        self.context_map[label] = self.context_map["latest"]
        return self

    def transform_context(
        self, label: str = "latest", func: Callable[[str], str] = lambda x: x
    ) -> "AIFlow":
        """
        Transform the context using a function.

        :param label: Label for the context
        :param func: Function to transform the context
        :return: self
        """
        if label != "latest" and label in self.context_map and func is not None:
            self.context_map[label] = func(self.context_map[label])
        return self

    def set_context_of(self, content: str = "", label: str = "latest") -> "AIFlow":
        """
        Set the context for a specified label.

        :param content: Content to set
        :param label: Label for the context
        :return: self
        """
        if label != "latest":
            self.context_map[label] = content
        return self

    def delete_context(self, label: str = "latest") -> "AIFlow":
        """
        Delete the context for a specified label.

        :param label: Label for the context
        :return: self
        """
        if label != "latest" and label in self.context_map:
            del self.context_map[label]
        return self

    def display_context_of(self, label: str = "latest") -> "AIFlow":
        """
        Show the context for a specified label.

        :param label: Label for the context
        :return: self
        """
        print(self.context_map[label])
        return self

    def display_context_keys(self) -> "AIFlow":
        """
        Show all context keys.

        :return: self
        """
        keys_list = list(self.context_map.keys())
        keys_str = ", ".join(keys_list)
        print(keys_str)
        return self

    def return_context_keys(self) -> List[str]:
        """
        Return all context keys.

        :return: List of context keys
        """
        return self.context_map.keys()

    def load_to_context(self, filename: str, label: str = "latest_file") -> "AIFlow":
        """
        Load content from a file into the context.

        :param filename: Name of the file to load content from
        :param label: Label for the context
        :return: self
        """
        try:
            with open(filename, "r") as file:
                content = file.read()
                self.context_map[label] = content
        except FileNotFoundError:
            print(f"The file {filename} does not exist.")
        except Exception as e:
            print(f"An error occurred: {e}")
        return self

    def save_context_to_file(
        self, label: str = "latest", filename: str = ""
    ) -> "AIFlow":
        """
        Dump the context to a file.

        :param label: Label for the context
        :param filename: Name of the file to dump content to
        :return: self
        """
        if self.default_folder_for_output != "":
            filename_2 = os.path.join(
                self.default_folder_for_output, f"context_{label}.txt"
            )
        else:
            filename_2 = f"context_{label}.txt"

        if filename != "":
            filename_2 = filename

        with open(filename_2, "w") as file:
            file.write(str(self.context_map[label]))

        return self

    def save_context_to_files(self) -> "AIFlow":
        """
        Dump all contexts to files.

        :return: self
        """
        for key, value in self.context_map.items():
            self.save_context_to_file(label=key)
        return self

    def load_multiple_context_from_file(self, filename: str = "") -> "AIFlow":
        """
        Load a text file and process its content grouped by keywords.

        Reads a file line by line, identifying sections where each keyword is
        followed by content. The keyword lines end with a colon, and all lines
        until the next keyword belong to that keyword's content. For each
        keyword-content block found, it calls `set_context_of` with the keyword
        as the label and the content as the section's text.

        :param filename: Name of the file to load and process
        :return: None
        """
        with open(filename, "r") as file:
            current_keyword = None
            current_content = []

            for line in file:
                line = line.strip()  # Remove whitespace and newline characters

                # Check if the line is a keyword (contains a colon at the end)
                if line.endswith(":"):
                    # Process the previous keyword-content block
                    if current_keyword:
                        # Join the collected lines and pass to the function
                        content = "\n".join(current_content)
                        self.set_context_of(label=current_keyword, content=content)

                        if self.verbose:
                            print(f"Stored context {current_keyword}")

                    # Set new keyword and reset content list
                    current_keyword = line[:-1]  # Remove the colon
                    current_content = []

                else:
                    # Accumulate lines under the current keyword
                    current_content.append(line)

            # Process the last keyword-content block if it exists
            if current_keyword:
                content = "\n".join(current_content)
                self.set_context_of(label=current_keyword, content=content)

                if self.verbose:
                    print(f"Stored context {current_keyword}")

            if self.save_state_per_step:
                self.save_internal_state()

        def save_context_to_markdown(
            self, output_filename: str = "content.md"
        ) -> "AIFlow":
            """
            Dump the context to a markdown file.

            :param output_filename: Name of the markdown file
            :return: self
            """
            with open(output_filename, "w") as file:
                for chapter, content in self.context_map.items():
                    file.write(f"# {chapter}\n\n")
                    file.write(f"{content}\n\n")
            return self

        def generate_heading_for_context(
            self, label: str, prompt: str, replace: bool
        ) -> None:
            """
            Generate a heading for a single context.

            :param label: Label for the context
            :param prompt: Prompt for generating the heading
            :param replace: Whether to replace the existing heading
            """
            content = self.context_map.get(label, "")
            if not content:
                return

            heading_label = label + "_heading"
            existing_heading = self.context_map.get(heading_label)

            if replace or not existing_heading:
                full_prompt = f"{prompt}{content}"
                messages = [{"role": "user", "content": full_prompt}]
                response = self.call_openai_chat_api(messages=messages)
                self.set_context_of(label=heading_label, content=response)

                if self.save_state_per_step:
                    self.save_internal_state()

    def generate_headings_for_contexts(
        self,
        labels: List[str] = [],
        prompt: str = "Generate a short 10 word summary of the following content:\n",
        replace: bool = True,
    ) -> "AIFlow":
        """
        Generate headings for multiple contexts.

        :param labels: List of labels for the contexts
        :param prompt: Prompt for generating the headings
        :param replace: Whether to replace the existing headings
        :return: self
        """
        for label in labels:
            self.generate_heading_for_context(label, prompt, replace)
        return self

    def save_context_to_docx(
        self, output_filename: str, chapters_to_include: List[str] = []
    ) -> "AIFlow":
        """
        Save the context to a DOCX file.

        :param output_filename: Name of the DOCX file
        :param chapters_to_include: List of chapters to include
        :return: self
        """
        document = Document()
        for chapter, content in self.context_map.items():
            if chapter in chapters_to_include:
                heading_key = chapter + "_heading"
                if heading_key in self.context_map:
                    document.add_heading(self.context_map[heading_key], level=1)
                else:
                    document.add_heading(chapter, level=1)

                document.add_paragraph(str(content))

        document.save(output_filename)

        return self

    def save_context_to_html(
        self, output_filename: str, chapters_to_include: List[str] = []
    ) -> "AIFlow":
        """
        Save the context to an HTML file.

        :param output_filename: Name of the HTML file
        :param chapters_to_include: List of chapters to include
        :return: self
        """
        html_content = "<html><body>"

        if chapters_to_include:
            for chapter in chapters_to_include:
                if chapter in self.context_map:
                    heading_key = chapter + "_heading"
                    if heading_key in self.context_map:
                        heading = self.context_map[heading_key]
                    else:
                        heading = chapter

                    html_content += f"<h1>{heading}</h1>"
                    html_content += markdown.markdown(str(self.context_map[chapter]))
        else:
            for chapter, content in self.context_map.items():
                heading_key = chapter + "_heading"
                if heading_key in self.context_map:
                    heading = self.context_map[heading_key]
                else:
                    heading = chapter

                html_content += f"<h1>{heading}</h1>"
                html_content += markdown.markdown(str(content))

        html_content += "</body></html>"

        # Ensure the directory exists
        os.makedirs(os.path.dirname(output_filename), exist_ok=True)

        with open(output_filename, "w", encoding="utf-8") as f:
            f.write(html_content)

        return self

    #
    # Functions to support inclusion in other chat instances. Only implemented by returning strings
    #
    def get_latest_context_as_text(self) -> str:
        """
        Get the latest context as text.

        :return: Latest context as text
        """
        return self.context_map["latest"]

    def get_context_as_text(self, label: str = "latest") -> str:
        """
        Get the context as text for a specified label.

        :param label: Label for the context
        :return: Context as text
        """
        return self.context_map[label]

    def get_context_as_json(self, label: str = "latest") -> str:
        """
        Get the context as text for a specified label as json object

        :param label: Label for the context
        :return: Context as json
        """
        return json.loads(self.context_map[label])

    def get_reduced_chat_messages_as_text(
        self, func: Optional[Callable[[List[Dict[str, str]]], str]]
    ) -> str:
        """
        Get reduced chat messages as text using a function.

        :param func: Function to reduce messages
        :return: Reduced chat messages as text
        """
        if func is not None:
            return func(self.chat_messages)
        return self

    def display_latest_context_as_markdown(self) -> None:
        """
        Display the latest context as markdown.

        :return: None
        """
        return display(Markdown(self.context_map["latest"]))

    def display_context_as_markdown(self, label: str = "latest") -> None:
        """
        Display the context as markdown for a specified label.

        :param label: Label for the context
        :return: None
        """
        return display(Markdown(self.context_map[label]))

    #
    # Image generation
    # dall-e-3 dall-e-2
    # 1024x1024 512x512
    # https://platform.openai.com/docs/api-reference/images/create
    #
    def generate_image(
        self,
        model: str = "dall-e-2",
        style: str = "vivid",
        response_format: str = "url",
        prompt: str = "A white siamese cat",
        size: str = "1024x1024",
        quality: str = "standard",
        n: int = 1,
        label: str = "latest_image",
        html: bool = False,
    ) -> "AIFlow":
        """
        Generate an image.

        :param model: Model to use for image generation
        :param style: Style of the image
        :param response_format: Format of the response (url or b64_json)
        :param prompt: Prompt for image generation
        :param size: Size of the image
        :param quality: Quality of the image
        :param n: Number of images to generate
        :param label: Label for the generated image
        :param html: Whether to return HTML for displaying the image
        :return: self
        """
        print(f"Generating image with prompt: {prompt}")
        response = self.client.images.generate(
            model=model,
            prompt=prompt,
            size=size,
            quality=quality,
            n=n,
            response_format=response_format,
            style=style,
        )

        image_data = response.data[0]
        self.images_map[label] = (
            image_data.url if response_format == "url" else image_data.b64_json
        )
        self.images_map["latest_image"] = self.images_map[label]

        self.context_map["latest_image_prompt"] = prompt
        self.context_map["latest_revised_image_prompt"] = image_data.revised_prompt

        if html:
            return HTML(f'<img src="{self.images_map[label]}" />')

        return self

    def save_image_to_file(
        self, label: str = "latest_image", filename: str = ""
    ) -> "AIFlow":
        """
        Save the generated image to a file.

        :param label: Label for the generated image
        :param filename: Name of the file to save the image
        :return: self
        """
        """
        Save the generated image to a file.

        :param label: Label for the generated image
        :param filename: Name of the file to save the image
        :return: self
        """
        output_filename = filename if filename else f"image_{label}.jpg"
        if not output_filename.lower().endswith(".jpg"):
            output_filename += ".jpg"

        if self.verbose:
            print(f"Saving {label} to {output_filename}")

        if label in self.images_map:
            urllib.request.urlretrieve(self.images_map[label], output_filename)

        return self

    #
    # Vision model
    # https://platform.openai.com/docs/guides/vision
    #
    def analyze_image(
        self,
        image: str = "",
        prompt: str = "What's in this image?",
        model: str = "gpt-4o",
        label: str = "latest",
        detail: str = "low",
        max_tokens: int = 300,
    ) -> "AIFlow":
        """
        Analyze an image.

        :param image: URL of the image
        :param prompt: Prompt for image analysis
        :param model: Model to use for analysis
        :param label: Label for the context
        :param detail: Detail level of the analysis
        :param max_tokens: Maximum tokens for the response
        :return: self
        """
        if image == "":
            return self

        response = self.client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": image, "detail": detail},
                        },
                    ],
                }
            ],
            max_tokens=max_tokens,
        )

        # increase the tokens using completion.usage
        self.update_token_usage(response)

        if self.verbose:
            print(response.choices[0].message)

        self.context_map["latest"] = response.choices[0].message
        self.context_map[label] = response.choices[0].message

        return self

    #
    # MP3 generation
    # https://platform.openai.com/docs/api-reference/audio/createSpeech
    #
    def generate_speech(
        self,
        model: str = "tts-1",
        voice: str = "alloy",
        response_format: str = "mp3",
        prompt: str = "A white siamese cat",
        speed: int = 1,
        filename: str = "",
        label: str = "latest_speech",
        html: bool = False,
    ) -> "AIFlow":
        """
        Generate speech from text.

        :param model: Model to use for speech generation
        :param voice: Voice to use
        :param response_format: Format of the response (mp3 or other)
        :param prompt: Prompt for speech generation
        :param speed: Speed of the speech
        :param filename: Name of the file to save the speech
        :param label: Label for the generated speech
        :param html: Whether to return HTML for downloading the speech
        :return: self
        """
        if filename == "":
            filename = label + ".mp3"

        speech_file_path = Path(__file__).parent / filename

        response = self.client.audio.speech.create(
            model=model,
            input=prompt,
            voice=voice,
            speed=speed,
            response_format=response_format,
        )
        response.stream_to_file(speech_file_path)

        # increase the tokens using completion.usage
        self.update_token_usage(response)

        self.audio_map[label] = filename
        self.context_map["latest_speech_file"] = filename
        self.context_map["latest_speech_prompt"] = prompt

        if html:
            return HTML(
                f'<a href="{speech_file_path}" download>Click here to download the speech generated</a>'
            )

        return self

    # Sound transcription
    # https://platform.openai.com/docs/api-reference/audio/createTranscription
    def transcribe_audio(
        self,
        filename: str = "",
        model: str = "whisper-1",
        language: str = "en",
        prompt: str = "",
        response_format: str = "text",
        temperature: int = 0,
        label: str = "latest",
    ) -> "AIFlow":
        """
        Transcribe audio to text.

        :param filename: Name of the audio file
        :param model: Model to use for transcription
        :param language: Language of the audio
        :param prompt: Prompt for transcription
        :param response_format: Format of the response (text or other)
        :param temperature: Temperature for the model
        :param label: Label for the context
        :return: self
        """
        if filename == "":
            return self

        audio_file = open(filename, "rb")
        transcript = self.client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            language=language,
            prompt=prompt,
            response_format=response_format,
            temperature=temperature,
        )

        self.context_map["latest"] = transcript
        self.context_map[label] = transcript

        print(transcript)
        return self

    #
    # Moderation
    # https://platform.openai.com/docs/guides/moderation/overview
    #
    def moderate_content(
        self, prompt: str = "", label: str = "latest_moderation"
    ) -> "AIFlow":
        """
        Create a moderation request.

        :param prompt: Prompt for moderation
        :param label: Label for the context
        :return: self
        """
        if prompt == "":
            return self

        moderation = self.client.moderations.create(input=prompt)

        if self.verbose:
            print(moderation.to_json())

        self.context_map["latest_moderation"] = moderation.to_json()
        self.context_map[label] = moderation.to_json()

        return self

    def update_token_usage(self, usage: Dict[str, int]) -> None:
        """
        Update the token usage statistics.

        :param usage: Dictionary with token usage
        :return: None
        """
        """
        Update the token usage statistics.

        :param usage: Dictionary with token usage
        :return: None
        """
        # increase the tokens using completion.usage
        self.completion_tokens += usage.completion_tokens
        self.prompt_tokens += usage.prompt_tokens
        self.total_tokens += usage.total_tokens

        self.context_map["_completion_tokens"] = self.completion_tokens
        self.context_map["_prompt_tokens"] = self.prompt_tokens
        self.context_map["_total_tokens"] = self.total_tokens

        return
